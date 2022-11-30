import email
from multiprocessing.sharedctypes import Value
from this import d
from django.shortcuts import render, redirect
from django.http import HttpResponseRedirect
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth.models import User
import sweetify 
import datetime 
from django.contrib.auth.decorators import login_required
from docx import Document
from django.http import HttpResponse
from django.http import FileResponse
import openpyxl
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, A2, A3
from reportlab.lib.pagesizes import mm
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER 
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, NextPageTemplate
import numpy as np
import pandas as pd

from django.contrib import messages
from src.forms import (
    FormClienteNormal1, FormClienteNormal2, FormClienteNormal3, addproductsForm, FormVendedorPersona,
    FormVendedorUsuario, FormVendedorEmpleado, FormEmpleadoPersona, FormEmpleadoUsuario, FormEmpleadoEmpleado,
    FormProveedor, FormProductoproveedor, FormBodega, FormClienteEmpresa, FormCliente, FormTipodocumento, FormVenta, 
    FormDocu, FormProveedorUsuario,FormDireProvee, FormFamiliaProduct, FormDetalleorden, FormOrdencompra, FormAccionpagina
)

from .models import (
    Detalleorden, Estadoorden, Guiadespacho, Ordencompra, Persona, Direccion, Usuario, Cliente, Estado, Comuna, Tipobarrio, Despacho,
    Tipovivienda, Rolusuario, Direccioncliente, Empresa, Proveedor, Tipoproducto, Producto, Familiaproducto, Tipopago,
    Empleado, Cargo, Tiporubro, Recepcion, Productoproveedor, Bodega, Boleta, Factura, Venta, Tipodocumento, Detalleventa, Boleta,
    Notacredito, Accionpagina
)

@login_required(login_url="ingreso")
def Index(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    context = {
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'index.html', context)

def Ingreso(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    if request.method == 'POST':

        username = request.POST.get('usuario')
        password = request.POST.get('contraseña')

        usuario_correo = Usuario.objects.filter(nombreusuario=username).exists()

        if usuario_correo == True:

            user = authenticate(request, username = username, password = password)

            if user is not None:

                login(request, user)
                sweetify.success(request, 'Ingreso realizado con exito')
                return redirect('index')

            else:

                sweetify.warning(request, 'Usuario y/o contraseña inválidos.')
        else:

            sweetify.warning(request, 'Usuario y/o contraseña inválidos.')

    context = {
        'tipo_usuario': tipo_usuario
    }
    
    return render(request, 'ingreso/ingreso_usuarios.html', context)

#****************************Creacion de archivos******************************************

def creacion_excel(nombre_archivo, lista, tipo_doc, extension):
    book = openpyxl.Workbook()  # Se crea un libro
    sheet = book.active  # Se activa la primera hojar
    sheet.title = f"{nombre_archivo}"
    cont = 0
    cont2 = 0

    for i in lista:
        cont += 1
        for j in i:
            
            cont2 += 1
            val = sheet.cell(row=cont, column=cont2)
            val.value = j
        cont2 = 0    

    cont = 0
    cont2 = 0
    
    response = HttpResponse(content_type=f"application/{tipo_doc}")
    contenido = "attachment; filename = {0}.{1}".format(nombre_archivo, extension)
    response["Content-Disposition"] = contenido
    book.save(response)

    return response

def creacion_pdf(lista,tipo_doc,tamaño_pagina, nombre, extension,lista2 = None, valor = None):
    response = HttpResponse(content_type=f'application/{tipo_doc}')  

    buff = BytesIO()  

    doc = SimpleDocTemplate(buff,  
        pagesize=tamaño_pagina,  
        rightMargin=40,  
        leftMargin=40,  
        topMargin=60,  
        bottomMargin=18,  
    ) 
    
    data = []  
    styles = getSampleStyleSheet()  
    styles = styles['Heading1']
    styles.alignment = TA_CENTER 

    header = Paragraph(f"{nombre}", styles)  
    
    data.append(header)  

    t = Table(lista)  

    t.setStyle(TableStyle(  
        [  
        ('GRID', (0, 0), (12, -1), 1, colors.dodgerblue),  
        ('LINEBELOW', (0, 0), (-1, 0), 3, colors.darkblue),  
        ('BACKGROUND', (0, 0), (-1, 0), colors.dodgerblue)  
        ]   
    ))  
    
    data.append(t)

    if lista2:

        styles = getSampleStyleSheet()  
        styles.pageBreakBefore = 2
        styles = styles['Heading1']
        styles.alignment = TA_CENTER 
        header = Paragraph(f" ", styles)  
        
        data.append(header)  

        t = Table(lista2)  

        t.setStyle(TableStyle(  
            [  
            ('GRID', (0, 0), (12, -1), 1, colors.dodgerblue),  
            ('LINEBELOW', (0, 0), (-1, 0), 3, colors.darkblue),  
            ('BACKGROUND', (0, 0), (-1, 0), colors.dodgerblue)  
            ]  
        ))  
        
        data.append(t)

    doc.build(data)  
    response.write(buff.getvalue())   

    buff.seek(0)

    respuesta = FileResponse(buff, as_attachment=valor, filename=f'{nombre}.{extension}')

    return respuesta

def creacion_doc(lista, nombre_archivo):
    document = Document()
    document.add_heading('Document Title', 0)

    filas = 0
    for x in lista:
        columnas = len(x)
        filas += 1

    # add grid table
    table = document.add_table(rows=filas, cols=columnas, style="Table Grid")

    # access first row's cells
    heading_row = table.rows[0].cells

    # add headings
    cont = 0
    for value in lista[0]:
        heading_row[cont].text = value
        cont += 1

    lista.pop(0)
    cont = 0
    cont2 = 0

    for value in lista:
        cont += 1
        data_row = table.rows[cont].cells

        for x in value:
            data_row[cont2].text = f'{x}'
            cont2 += 1
        cont2 = 0     

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={nombre_archivo}.docx'
    document.save(response)

    return response

#************************************Informes**************************************

@login_required(login_url="ingreso")
def informe_productos(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None
    
    form1 = addproductsForm(request.POST)

    if request.method == 'POST':
        vistaPrevia = request.POST.get('vistaPrevia')
        descargarInforme = request.POST.get('descargarInforme')
        productos = request.POST.get('productos')
        precio = request.POST.get('precio')
        nombre = request.POST.get('nombre')
        precioCompra = request.POST.get('precioCompra')
        stockCritico = request.POST.get('stockCritico')
        fechaVencimiento = request.POST.get('fechaVencimiento')
        codigoBarra = request.POST.get('codigoBarra')
        stock_producto = request.POST.get('stock')
        stockCheck = request.POST.get('stockCheck')
        estado = request.POST.get('estado')
        estadoCheck = request.POST.get('estadoCheck')
        familiaProducto = request.POST.get('familiaProducto')
        familiaProductoCheck = request.POST.get('familiaProductoCheck')
        nomFamiliaProducto = request.POST.get('familia_producto')
        tipoProducto = request.POST.get('tipoProducto')
        nomTipoProducto = request.POST.get('tipo_producto')
        tipoProductoCheck = request.POST.get('tipoProductoCheck')
        familiaproid = request.POST.get('familiaproid')
        tipoproductoid = request.POST.get('tipoproductoid')

        lista = []
        visitas = []

        tipoInforme = request.POST.get('informeCheck')
        
        np_array = []
        np_array = np.array(lista)

        if productos == "on":
            productos = Producto.objects.filter().values_list("nombre","precio","stock","stockcritico","fechavencimiento","codigo","familiaproid__descripcion","tipoproductoid__descripcion","estadoid__descripcion","bodegaid__pasillo","bodegaid__estante","bodegaid__casillero").order_by("productoid")
            if stockCheck == "conStock":
                productos = productos.filter(stock__gt=0)
            if stockCheck == "sinStock":
                productos = productos.filter(stock__lt=0)
            if estadoCheck == "disponible":
                productos = productos.filter(estadoid__estadoid=1)
            if estadoCheck == "noDisponible":
                productos = productos.filter(estadoid__estadoid=2)
            if familiaProductoCheck == "porNombreF":
                productos = productos.filter(familiaproid__familiaproid = familiaproid)
            if tipoProductoCheck == "porNombreT":
                productos = productos.filter(tipoproductoid__tipoproductoid = tipoproductoid)
            columnas = (["Nombre","Precio","Stock","Stock Critico","Fecha Vencimiento","Codigo","Familia Producto","Tipo Producto","Estado","Pasillo","Estante","Casillero"])
 
            np_array = np.array(productos)
            df = pd.DataFrame(np_array, columns = columnas)
            df["Bodega"] = "Pasillo:" +df["Pasillo"] + " Estante:" + df["Estante"]+ " Casillero:" + df["Casillero"]
            df = df.drop(['Pasillo'], axis=1)
            df = df.drop(['Estante'], axis=1)
            df = df.drop(['Casillero'], axis=1)

            if nombre == None:
                df = df.drop(['Nombre'], axis=1)
            if precio == None:
                df = df.drop(['Precio'], axis=1)
            if stock_producto == None:
                df = df.drop(['Stock'], axis=1)
            if stockCritico == None:
                df = df.drop(['Stock Critico'], axis=1)
            if estado == None:
                df = df.drop(['Estado'], axis=1)
            if fechaVencimiento == None:
                df = df.drop(['Fecha Vencimiento'], axis=1)
            if codigoBarra == None:
                df = df.drop(['Codigo'], axis=1)
            if codigoBarra == None:
                df = df.drop(['Bodega'], axis=1)
            if familiaProducto == None:
                df = df.drop(['Familia Producto'], axis=1)
            if tipoProducto == None:
                df = df.drop(['Tipo Producto'], axis=1)

            if df.shape[1] > 6:
                length_dataframe = df.shape[1]
                df2= df.iloc[:, 6:int(length_dataframe)] 
                df= df.iloc[:, 0:6]

                lista_productos2 = df2.values.tolist() 
                columnas_df2 = df2.columns.values.tolist() 
                lista_productos2.insert(0, columnas_df2)

                lista_productos = df.values.tolist() 
                columnas_df = df.columns.values.tolist() 
                lista_productos.insert(0, columnas_df)
            else:
                lista_productos = df.values.tolist() 
                columnas_df = df.columns.values.tolist() 
                lista_productos.insert(0, columnas_df)

            if tipoInforme == "informeExcel":
                
                nombre_archivo = "Productos"
                tipo_doc = 'ms-excel'
                extension = 'xlsx'
                
                return  creacion_excel(nombre_archivo, lista, tipo_doc, extension)
                
                    

            if tipoInforme == "informePdf":

                tipo_doc = 'pdf'
                extension = 'pdf'
                nombre = 'Informe Productos'
                if vistaPrevia:
                    if df.shape[1] > 6:
                        return creacion_pdf(lista_productos,tipo_doc,A4,nombre,extension, lista_productos2, valor=False)
                    else:
                        return creacion_pdf(lista_productos,tipo_doc,A4,nombre,extension,valor=False)
                else:
                    if df.shape[1] > 6:
                        return creacion_pdf(lista_productos,tipo_doc,A4,nombre,extension, lista_productos2, valor=True)
                    else:
                        return creacion_pdf(lista_productos,tipo_doc,A4,nombre,extension,valor=True)

            if tipoInforme == "informeWord": 

                tipo_doc = 'ms-word'
                extension = 'docx'
                
                nombre = 'Productos'
                if vistaPrevia:
                    if visitas == []:
                        
                        return creacion_doc(lista_productos,nombre)
                    else:
                        return creacion_doc(lista_productos,nombre)
                else:
                    if visitas == []:
                        return creacion_doc(lista_productos,nombre)
                    else:
                        return creacion_doc(lista_productos,nombre)

    context = {
        'form1': form1,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'informes/informe_productos.html', context)

@login_required(login_url="ingreso")
def informe_proveedores(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None
    
    form1 = FormProveedor(request.POST)

    if request.method == 'POST':
        vistaPrevia = request.POST.get('vistaPrevia')
        proveedores = request.POST.get('proveedores')
        correoP = request.POST.get('correoP')
        telefonoP = request.POST.get('telefonoP')
        direccionP = request.POST.get('direccionP')
        estadoP = request.POST.get('estadoP')
        estadoProveedorCheck = request.POST.get('estadoProveedorCheck')
        nomProveedor = request.POST.get('proveedor')
        categoriaP = request.POST.get('categoriaP')
        categoriaPCheck = request.POST.get('categoriaPCheck')
        nomCategoria = request.POST.get('categoria_proveedor')
        rubroid = request.POST.get('rubroid')

        lista = []
        visitas = []

        tipoInforme = request.POST.get('informeCheck')
        
        np_array = []
        np_array = np.array(lista)

        if proveedores == "on":
            proveedores = Proveedor.objects.all().values_list("razonsocial","rutcuerpo","dv","fono","rubroid__descripcion","direccionid__calle","direccionid__numero","direccionid__comunaid__nombre","estadoid__descripcion").order_by("proveedorid")
            columnas = (["Razon Social","RUN", "DV", "Telefono", "Rubro","calle","numero","comuna","Estado"])

            if estadoProveedorCheck == "activosP":
                proveedores = proveedores.filter(estadoid__estadoid=1)
            if estadoProveedorCheck == "bloqueadosP":
                proveedores = proveedores.filter(estadoid__estadoid=2)
            if categoriaPCheck == "porcategoriaP":
                proveedores = proveedores.filter(rubroid__rubroid = rubroid)
                
            np_array = np.array(proveedores)
            df = pd.DataFrame(np_array, columns = columnas)

            df["Direccion"] = df["calle"] + " " + df["numero"] + "," + df["comuna"]
            df = df.drop(['calle'], axis=1)
            df = df.drop(['numero'], axis=1)
            df = df.drop(['comuna'], axis=1)

            if telefonoP == None:
                df = df.drop(['Telefono'], axis=1)
            if direccionP == None:
                df = df.drop(['Direccion'], axis=1)

            if df.shape[1] > 6:
                length_dataframe = df.shape[1]
                df2 = df.iloc[:, 6:int(length_dataframe)] 
                df = df.iloc[:, 0:6]

                lista_proveedores2 = df2.values.tolist() 
                columnas_df2 = df2.columns.values.tolist() 
                lista_proveedores2.insert(0, columnas_df2)

                lista_proveedores = df.values.tolist() 
                columnas_df = df.columns.values.tolist() 
                lista_proveedores.insert(0, columnas_df)
            else:
                lista_proveedores = df.values.tolist() 
                columnas_df = df.columns.values.tolist() 
                lista_proveedores.insert(0, columnas_df)

            if tipoInforme == "informeExcel":
                
                nombre_archivo = "Proveedores"
                tipo_doc = 'ms-excel'
                extension = 'xlsx'

                return  creacion_excel(nombre_archivo, lista_proveedores, tipo_doc, extension)
                
                
            if tipoInforme == "informePdf":

                tipo_doc = 'pdf'
                extension = 'pdf'
                nombre = 'Informe Proveedores'
                if vistaPrevia:
                    if df.shape[1] >= 6:
                        return creacion_pdf(lista_proveedores,tipo_doc,A4,nombre,extension, lista_proveedores2, valor=False)
                    else:
                        return creacion_pdf(lista_proveedores,tipo_doc,A4,nombre,extension,valor=False)
                else:
                    if df.shape[1] >= 6:
                        return creacion_pdf(lista_proveedores,tipo_doc,A4,nombre,extension, lista_proveedores2, valor=True)
                    else:
                        return creacion_pdf(lista_proveedores,tipo_doc,A4,nombre,extension,valor=True)

            if tipoInforme == "informeWord": 

                tipo_doc = 'ms-word'
                extension = 'docx'
                
                nombre = 'Proveedores'
                if vistaPrevia:
                    if visitas == []:
                        return creacion_doc(lista_proveedores,nombre)
                    else:
                        return creacion_doc(lista_proveedores,nombre)
                else:
                    if visitas == []:
                        return creacion_doc(lista_proveedores,nombre)
                    else:
                        return creacion_doc(lista_proveedores,nombre)

    context = {
        'form1':form1,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'informes/informe_proveedores.html', context)

@login_required(login_url="ingreso")
def Seleccion_informe(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None
    
    context = {
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'informes/seleccion_informe.html', context)

@login_required(login_url="ingreso")
def informe_pedidos(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    form1 = FormOrdencompra(request.POST)

    if request.method == 'POST':
        ordenes = request.POST.get('ordenes')
        fechaOrdenP = request.POST.get('fechaOrdenP')
        fechaLlegadaOrdenP = request.POST.get('fechaLlegadaOrdenP')
        fechaRecepOrdenP = request.POST.get('fechaRecepOrdenP')
        horaRecepOrdenP = request.POST.get('horaRecepOrdenP')
        cantidadOrdenP = request.POST.get('cantidadOrdenP')
        productoOrdenP = request.POST.get('productoOrdenP')
        estadoordenid = request.POST.get('estadoordenid')
        proveedorid = request.POST.get('proveedorid')

        estadoRecepcionO = request.POST.get('estadoRecepcionO')
        razonSocialOrdenP = request.POST.get('razonSocialOrdenP')
        razonSocialOrdenPCheck = request.POST.get('razonSocialOrdenPCheck')
        nomProveedor = request.POST.get('proveedor')
        tipoInforme = request.POST.get('informeCheck')
        vistaPrevia = request.POST.get('vistaPrevia')
        visitasPagina = request.POST.get('visitas')
        
        lista = []
        visitas = []

        np_array = []
        np_array = np.array(lista)

        if ordenes == "on":
            ordenes = Detalleorden.objects.all().values_list("ordenid__proveedorid__razonsocial","ordenid__fechapedido","ordenid__estadoordenid__descripcion","productoid__nombre","cantidad").order_by("detalleid")
            columnas = (["Razon Social","Fecha Pedido", "Estado", "Producto", "Cantidad"])
            if estadoRecepcionO == "on":   
                ordenes = ordenes.filter(ordenid__estadoordenid = estadoordenid)
            if razonSocialOrdenPCheck == "porRazonSocialOrdenP":
                ordenes = ordenes.filter(ordenid__proveedorid = proveedorid)

            if len(ordenes) < 1:
                sweetify.warning(request, 'No existen registros con los parámetros seleccionados')
                return redirect('informe_pedidos')
            else:
                np_array = np.array(ordenes)
                df = pd.DataFrame(np_array, columns = columnas)

                if cantidadOrdenP == None:
                    df = df.drop(['Cantidad'], axis=1)
                if productoOrdenP == None:
                    df = df.drop(['Producto'], axis=1)

                lista_proveedores = df.values.tolist() 
                columnas_df = df.columns.values.tolist() 
                lista_proveedores.insert(0, columnas_df)
            
            if tipoInforme == "informeExcel":
                nombre_archivo = "Pedidos"
                tipo_doc = 'ms-excel'
                extension = 'xlsx'

                creacion_excel(nombre_archivo, lista_proveedores, tipo_doc, extension)
                

            if tipoInforme == "informePdf":
                tipo_doc = 'pdf'
                extension = 'pdf'
                nombre = 'Informe Pedidos'
                if vistaPrevia:
                    return creacion_pdf(lista_proveedores,tipo_doc,A4,nombre,extension, valor=False)
                else:
                    return creacion_pdf(lista_proveedores,tipo_doc,A4,nombre,extension, valor=True)

            if tipoInforme == "informeWord": 
                tipo_doc = 'ms-word'
                extension = 'docx'
                nombre = 'Pedidos'
                if vistaPrevia:
                    if visitas == []:
                        return creacion_doc(lista_proveedores,nombre)
                    else:
                        return creacion_doc(lista_proveedores,nombre)
                else:
                    if visitas == []:
                        return creacion_doc(lista_proveedores,nombre)
                    else:
                        return creacion_doc(lista_proveedores,nombre)

    context = {
        'form1':form1,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'informes/informe_pedidos.html', context)

@login_required(login_url="ingreso")
def informe_ventas(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    form1 = FormVenta(request.POST)
    
    if request.method == 'POST':

        ventas = request.POST.get('ventas')
        fechaVenta = request.POST.get('fechaVenta')
        montoTotal = request.POST.get('montoTotal')
        nombreCliente = request.POST.get('nombreCliente')
        nombreProducto = request.POST.get('nombreProducto')
        cantidadVenta = request.POST.get('cantidadVenta')
        subTotal = request.POST.get('subTotal')

        DocuTributario = request.POST.get('DocuTributario')
        docuTributarioCheck = request.POST.get('docuTributarioCheck')
        tipodocumentoid = request.POST.get('tipodocumentoid')
        tipoPago = request.POST.get('tipoPago')
        TipoPagoCheck = request.POST.get('TipoPagoCheck')
        tipopagoid = request.POST.get('tipopagoid')

        tipoInforme = request.POST.get('informeCheck')
        vistaPrevia = request.POST.get('vistaPrevia')
        
        lista = []
        visitas = []

        tipoInforme = request.POST.get('informeCheck')
        
        np_array = []
        np_array = np.array(lista)

        if ventas == "on":
            ventas = Detalleventa.objects.all().values_list("nroventa","nroventa__fechaventa","nroventa__totalventa","nroventa__clienteid__personaid__nombres","nroventa__clienteid__empresaid__razonsocial", "productoid__nombre","cantidad","subtotal","nroventa__tipodocumentoid__descripcion","nroventa__tipopagoid__descripcion").order_by("nroventa")

            columnas = (["Nro Venta","Fecha Venta", "Valor Total", "Cliente", "Empresa", "Nombre Producto", "Cantidad", "Sub Total", "Documento Venta", "Tipo de Pago"])
            
            if docuTributarioCheck == "porDocuTributario":   
                ventas = ventas.filter(nroventa__tipodocumentoid = tipodocumentoid)
            if TipoPagoCheck == "porTipoPago":
                ventas = ventas.filter(nroventa__tipopagoid = tipopagoid)

            if len(ventas) < 1:
                sweetify.warning(request, 'No existen registros con los parámetros seleccionados')
                return redirect('informe_ventas')
            else:
                np_array = np.array(ventas)
                df = pd.DataFrame(np_array, columns = columnas)

                df["Nombre Cliente"] = df["Cliente"] + df["Empresa"]
                df = df.drop(['Cliente'], axis=1)
                df = df.drop(['Empresa'], axis=1)

                if fechaVenta == None:
                    df = df.drop(['Fecha Venta'], axis=1)
                if montoTotal == None:
                    df = df.drop(['Valor Total'], axis=1)
                if nombreCliente == None:
                    df = df.drop(['Nombre Cliente'], axis=1)
                if nombreProducto == None:
                    df = df.drop(['Nombre Producto'], axis=1)
                if cantidadVenta == None:
                    df = df.drop(['Cantidad'], axis=1)
                if subTotal == None:
                    df = df.drop(['Sub Total'], axis=1)

                if DocuTributario == None:
                    df = df.drop(['Documento Venta'], axis=1)
                if tipoPago == None:
                    df = df.drop(['Tipo de Pago'], axis=1)

            if df.shape[1] > 5:
                length_dataframe = df.shape[1]
                df2= df.iloc[:, 6:int(length_dataframe)] 
                df= df.iloc[:, 0:6]

                lista_ventas2 = df2.values.tolist() 
                columnas_df2 = df2.columns.values.tolist() 
                lista_ventas2.insert(0, columnas_df2)

                lista_ventas = df.values.tolist() 
                columnas_df = df.columns.values.tolist() 
                lista_ventas.insert(0, columnas_df)
            else:
                lista_ventas = df.values.tolist() 
                columnas_df = df.columns.values.tolist() 
                lista_ventas.insert(0, columnas_df)

            if tipoInforme == "informeExcel":
                nombre_archivo = "Informe Detallado Ventas"
                tipo_doc = 'ms-excel'
                extension = 'xlsx'

                return  creacion_excel(nombre_archivo, lista_ventas, tipo_doc, extension)
                

            if tipoInforme == "informePdf":
                tipo_doc = 'pdf'
                extension = 'pdf'
                nombre = 'Informe Detallado Ventas'
                if vistaPrevia:
                    if df.shape[1] > 5:
                        return creacion_pdf(lista_ventas,tipo_doc,A4,nombre,extension, lista_ventas2, valor=False)
                    else:
                        return creacion_pdf(lista_ventas,tipo_doc,A4,nombre,extension,valor=False)
                else:
                    if df.shape[1] > 5:
                        return creacion_pdf(lista_ventas,tipo_doc,A4,nombre,extension, lista_ventas2, valor=True)
                    else:
                        return creacion_pdf(lista_ventas,tipo_doc,A4,nombre,extension,valor=True)

            if tipoInforme == "informeWord": 
                tipo_doc = 'ms-word'
                extension = 'docx'
                nombre = 'Informe Detallado Ventas'
                if vistaPrevia:
                    if visitas == []:
                        return creacion_doc(lista_ventas,nombre)
                    else:
                        return creacion_doc(lista_ventas,nombre)
                else:
                    if visitas == []:
                        return creacion_doc(lista_ventas,nombre)
                    else:
                        return creacion_doc(lista_ventas,nombre)

    context = {
        'form1':form1,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'informes/informe_ventas.html', context)

@login_required(login_url="ingreso")
def informe_visitas(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    form1 = FormAccionpagina(request.POST)

    if request.method == 'POST':

        visitas = request.POST.get('visitas')
        fechaIn = request.POST.get('fechaIn')
        nombreModulo = request.POST.get('nombreModulo')
        fechain = request.POST.get('fechain')
        fechaInicio = request.POST.get('fechaInicio')
        porModuloCheck = request.POST.get('porModuloCheck')
        moduloCheck = request.POST.get('moduloCheck')
        detalleUsuario = request.POST.get('detalleUsuario')
        nombreUsuario = request.POST.get('nombreUsuario')
        correoUsuario = request.POST.get('correoUsuario')
        rolUsuario = request.POST.get('rolUsuario')

        tipoInforme = request.POST.get('informeCheck')
        vistaPrevia = request.POST.get('vistaPrevia')
        
        lista = []

        np_array = []
        np_array = np.array(lista)

        if visitas == "on":
            visitas = Accionpagina.objects.all().values_list("modulo","fechain","usuarioid__empresaid__razonsocial","usuarioid__personaid__nombres","usuarioid__email","usuarioid__rolid__descripcion").order_by("-fechain")
            
            columnas = (["Modulo", "Fecha Registro", "Empresa", "Persona", "Correo","Cargo"])

            if porModuloCheck == "porModulo":   
                visitas = visitas.filter(modulo__icontains = moduloCheck)

            if fechaInicio == "on":
                visitas = visitas.filter(fechain__gte = fechain)

            if len(visitas) < 1:
                sweetify.warning(request, 'No existen registros con los parámetros seleccionados')
                return redirect('informe_visitas')
            else:

                np_array = np.array(visitas)
                df = pd.DataFrame(np_array, columns = columnas)

                df["Usuario"] = df["Empresa"] + df["Persona"] 
                df = df.drop(['Empresa'], axis=1)
                df = df.drop(['Persona'], axis=1)

                df['Fecha Registro'] = df['Fecha Registro'].astype(str).str.slice(0, 19)

                if fechaIn == None:
                    df = df.drop(['Fecha Registro'], axis=1)
                if nombreModulo == None:
                    df = df.drop(['Modulo'], axis=1)
                if nombreUsuario == None:
                    df = df.drop(['Usuario'], axis=1)
                if correoUsuario == None:
                    df = df.drop(['Correo'], axis=1)
                if rolUsuario == None:
                    df = df.drop(['Cargo'], axis=1)
                
                if df.shape[1] > 6:
                    length_dataframe = df.shape[1]
                    df2 = df.iloc[:, 6:int(length_dataframe)] 
                    df = df.iloc[:, 0:6]

                    lista_visitas2 = df2.values.tolist() 
                    columnas_df2 = df2.columns.values.tolist() 
                    lista_visitas2.insert(0, columnas_df2)

                    lista_visitas = df.values.tolist() 
                    columnas_df = df.columns.values.tolist() 
                    lista_visitas.insert(0, columnas_df)
                else:
                    lista_visitas = df.values.tolist() 
                    columnas_df = df.columns.values.tolist() 
                    lista_visitas.insert(0, columnas_df)
                
            if tipoInforme == "informeExcel":
                nombre_archivo = "Pedidos"
                tipo_doc = 'ms-excel'
                extension = 'xlsx'
                return  creacion_excel(nombre_archivo, lista_visitas, tipo_doc, extension)
                

            if tipoInforme == "informePdf":
                tipo_doc = 'pdf'
                extension = 'pdf'
                nombre = 'Informe Pedidos'
                if vistaPrevia:
                    if df.shape[1] > 6:
                        return creacion_pdf(lista_visitas,tipo_doc,A4,nombre,extension, lista_visitas2, valor=False)
                    else:
                        return creacion_pdf(lista_visitas,tipo_doc,A4,nombre,extension,valor=False)
                else:
                    if df.shape[1] > 6:
                        return creacion_pdf(lista_visitas,tipo_doc,A4,nombre,extension, lista_visitas2, valor=True)
                    else:
                        return creacion_pdf(lista_visitas,tipo_doc,A4,nombre,extension,valor=True)

            if tipoInforme == "informeWord": 
                tipo_doc = 'ms-word'
                extension = 'docx'
                nombre = 'Pedidos'
                if vistaPrevia:
                    if visitas == []:
                        return creacion_doc(lista_visitas,nombre)
                    else:
                        return creacion_doc(lista_visitas,nombre)
                else:
                    if visitas == []:
                        return creacion_doc(lista_visitas,nombre)
                    else:
                        return creacion_doc(lista_visitas,nombre)

    context = {
        'form1':form1,
        'tipo_usuario': tipo_usuario,
    }

    return render(request, 'informes/informe_visitas.html', context)


def dashboard(request):

    if request.POST.get('VerPerfil') is not None:
        request.session['_ver_perfil'] = request.POST
        return redirect('ver_perfil')

    if Usuario.objects.filter(nombreusuario=request.user).exists():
        tipo_usuario = Usuario.objects.get(nombreusuario=request.user)
        tipo_usuario = tipo_usuario.rolid.descripcion
    else: 
        tipo_usuario = None

    # 1 Boleta
    # 2 Factura
    ventas = Venta.objects.all()
    boletas = 0
    facturas = 0
    for venta in ventas:
        if venta.tipodocumentoid.tipodocumentoid == 1:
            boletas+=1
        else:
            facturas+=1
    ventas_tipodocumento = [boletas, facturas]

    ventas = Venta.objects.all()
    despachos = Despacho.objects.all()

    ventas_total = 0
    despachos_total = 0
    for venta in ventas:
        ventas_total+=1

    for despacho in despachos:
        despachos_total+=1
        
    total_ventas_despachos = [ventas_total-despachos_total,despachos_total]
    import json

    detalle_ventas = Detalleventa.objects.all()
    productos_vendidos = []
    for detalle in detalle_ventas:
        n_producto = str(detalle.productoid)
        c_producto = int(detalle.cantidad)
        if len(productos_vendidos) == 0:
            productos_vendidos.append({"nombre": n_producto, "cantidad":c_producto})
        else:
            producto_found = next((product for product in productos_vendidos if product["nombre"] == n_producto), None)
            if producto_found:
                producto_found["cantidad"] = producto_found["cantidad"] +c_producto
            else:
                productos_vendidos.append({"nombre": n_producto, "cantidad":c_producto})

    nombre_productos = []
    cantidad_productos = []
    for producto in productos_vendidos:
        nombre_productos.append(producto["nombre"]) 
        cantidad_productos.append(producto["cantidad"]) 
    productos_vendidos = []
    productos_vendidos.append(nombre_productos)
    productos_vendidos.append(cantidad_productos)
    productos_vendidos = json.dumps(productos_vendidos)

    productos_stock = Producto.objects.all().order_by('stock')[:10]
    productos_peligro_stock = [] 
    nombre = []
    stock = []
    for pro_stock in productos_stock:
        nombre.append(str(pro_stock.nombre))
        stock.append(int(pro_stock.stock)-int(pro_stock.stockcritico))

    productos_peligro_stock.append(nombre)
    productos_peligro_stock.append(stock)
    productos_peligro_stock = json.dumps(productos_peligro_stock)
    
    context = {
        'ventasXDocumento': ventas_tipodocumento,
        'ventasXDespacho': total_ventas_despachos,
        'productosXcantidad': productos_vendidos,
        'totalVentas':ventas_total,
        'productoStock':productos_peligro_stock,
        'tipo_usuario': tipo_usuario,

    }

    return render(request, 'dashboard.html', context)